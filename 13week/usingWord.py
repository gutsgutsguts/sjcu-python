from docx import Document
from docx.shared import Cm, Inches

doc = Document()

doc.add_heading('파이썬 테스트입니다 - H1', level=1)
doc.add_paragraph('안녕하세요. 파이썬입니다.')
p = doc.add_paragraph('속성을 설정합니다.')
p.add_run('굵은글자입니다').bold = True

records = (
    (1, 'chris','30'),
    (2, 'tommy', '35'),
    (3, 'harry', '29')
)

table = doc.add_table(rows=1, cols=3)

table.style = doc.styles['Table Grid']

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Age'

for no, name, age in records:
   row_cells = table.add_row().cells
   row_cells[0].text = str(no) 
   row_cells[1].text = name 
   row_cells[2].text = age
   
   
doc.add_paragraph('')
doc.add_picture('./data/lena.jpg', width=Cm(16), height=Cm(9))


doc.save('test.docx')


doc = Document('test.docx')

for i, paragraph in enumerate(doc.paragraphs):
    print(str(i+1) + ": " + paragraph.text)
    
table = doc.tables[0]
for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            print(para.text)

for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            if(para.text == "chris"):
                para.add_run(' <-- me !')
                
doc.save('test_modify.docx')
                